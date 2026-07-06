"""The actual 'hands' of the team — real capabilities. FREE stack (Groq + DuckDuckGo)."""
import os
import time
import smtplib
import imaplib
import email
from email.message import EmailMessage
from email.header import decode_header

import httpx
from openai import OpenAI

from app.config import config

# Groq is OpenAI-compatible
llm = OpenAI(api_key=config.GROQ_API_KEY, base_url=config.GROQ_BASE_URL)


def _stamp(name, ext):
    safe = "".join(c for c in name if c.isalnum() or c in " -_")[:40].strip() or "file"
    return os.path.join(config.OUTPUT_DIR, f"{safe}-{int(time.time())}.{ext}")


# ---------------- RESEARCH (free web search) ----------------
def web_research(query: str) -> str:
    """Free web research: DuckDuckGo search -> summarize with the LLM."""
    snippets = []
    try:
        from ddgs import DDGS
        with DDGS() as ddg:
            for r in ddg.text(query, max_results=6):
                snippets.append(f"- {r.get('title','')}: {r.get('body','')} ({r.get('href','')})")
    except Exception as e:
        return f"[search error] {e}"

    if not snippets:
        return "(no search results)"

    context = "\n".join(snippets)
    try:
        resp = llm.chat.completions.create(
            model=config.MODEL, max_tokens=1200,
            messages=[
                {"role": "system", "content": "Summarize the search results into a clear, useful answer in simple language. Keep the source links."},
                {"role": "user", "content": f"Question: {query}\n\nSearch results:\n{context}"},
            ],
        )
        return resp.choices[0].message.content.strip()
    except Exception as e:
        return "Search results:\n" + context + f"\n[summary error: {e}]"


# ---------------- DOCUMENTS ----------------
def make_pdf(title: str, body: str) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    path = _stamp(title, "pdf")
    doc = SimpleDocTemplate(path, pagesize=A4)
    styles = getSampleStyleSheet()
    story = [Paragraph(title, styles["Title"]), Spacer(1, 16)]
    for para in body.split("\n"):
        if para.strip():
            story.append(Paragraph(para.strip(), styles["Normal"]))
            story.append(Spacer(1, 8))
    doc.build(story)
    return path


def make_docx(title: str, body: str) -> str:
    from docx import Document

    path = _stamp(title, "docx")
    d = Document()
    d.add_heading(title, level=0)
    for para in body.split("\n"):
        if para.strip():
            d.add_paragraph(para.strip())
    d.save(path)
    return path


def make_xlsx(title: str, rows: list) -> str:
    """rows = list of lists (first row = headers)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    path = _stamp(title, "xlsx")
    wb = Workbook()
    ws = wb.active
    ws.title = (title[:28] or "Sheet")
    for i, row in enumerate(rows):
        ws.append(row)
        if i == 0:
            for cell in ws[1]:
                cell.font = Font(bold=True)
    wb.save(path)
    return path


# ---------------- EMAIL ----------------
def read_inbox(n: int = 5) -> str:
    if not config.EMAIL_ADDRESS:
        return "[email not configured]"
    try:
        M = imaplib.IMAP4_SSL(config.IMAP_HOST)
        M.login(config.EMAIL_ADDRESS, config.EMAIL_APP_PASSWORD)
        M.select("INBOX")
        _, data = M.search(None, "ALL")
        ids = data[0].split()[-n:]
        out = []
        for i in reversed(ids):
            _, msg_data = M.fetch(i, "(RFC822)")
            msg = email.message_from_bytes(msg_data[0][1])
            subj = decode_header(msg.get("Subject", ""))[0][0]
            if isinstance(subj, bytes):
                subj = subj.decode(errors="ignore")
            out.append(f"From: {msg.get('From','')} | Subject: {subj}")
        M.logout()
        return "\n".join(out) or "(inbox empty)"
    except Exception as e:
        return f"[inbox error] {e}"


def send_email(to: str, subject: str, body: str) -> str:
    if not config.EMAIL_ADDRESS:
        return "[email not configured]"
    try:
        msg = EmailMessage()
        msg["From"] = config.EMAIL_ADDRESS
        msg["To"] = to
        msg["Subject"] = subject
        msg.set_content(body)
        with smtplib.SMTP_SSL(config.SMTP_HOST, config.SMTP_PORT) as s:
            s.login(config.EMAIL_ADDRESS, config.EMAIL_APP_PASSWORD)
            s.send_message(msg)
        return f"email sent to {to}"
    except Exception as e:
        return f"[send error] {e}"


# ---------------- WHATSAPP (Business Cloud API) ----------------
def whatsapp_send(to: str, text: str) -> str:
    if not config.WHATSAPP_TOKEN:
        return "[whatsapp not configured]"
    url = f"https://graph.facebook.com/v20.0/{config.WHATSAPP_PHONE_ID}/messages"
    try:
        r = httpx.post(
            url,
            headers={"Authorization": f"Bearer {config.WHATSAPP_TOKEN}"},
            json={"messaging_product": "whatsapp", "to": to,
                  "type": "text", "text": {"body": text}},
            timeout=30,
        )
        return f"whatsapp -> {to}: {r.status_code}"
    except Exception as e:
        return f"[whatsapp error] {e}"


# ---------------- BROWSER (human-jaisa Chrome) ----------------
def browse(url: str, instruction: str = "") -> str:
    """Open a page in headless Chrome and return visible text."""
    try:
        from playwright.sync_api import sync_playwright
        with sync_playwright() as p:
            browser = p.chromium.launch(headless=True, args=["--no-sandbox"])
            page = browser.new_page()
            page.goto(url, timeout=45000)
            page.wait_for_timeout(1500)
            text = page.inner_text("body")
            browser.close()
        return text[:6000]
    except Exception as e:
        return f"[browser error] {e}"
